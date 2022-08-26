from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import matplotlib.pyplot as plt
import numpy as np

data_f = open('report.csv') # CSV파일 읽기

scores = []

categories = []

for line in data_f:
    (name, hp) = line.split(',')
    scores.append(int(hp))
    categories.append(name)

data_f.close()

fig = plt.figure()
#ax = fig.add_subplot(1,1,1)

plt.plot(categories, scores, color = "gray", marker = "o")

plt.ylabel("Score")
plt.xlabel("Test")

fig = plt.gcf()
fig.savefig("Test.png")   # 선 그래프 생성

fig = plt.figure()

trash = [7,10,5,8,6,9,11]
plt.barh(categories, scores, color = 'black', alpha = 0.8)
plt.barh(categories, trash, left = scores, color = "gray")

fig = plt.gcf()
fig.savefig("Score.png") # 누적 가로 막대 그래프 생성

scores = [*scores, scores[0]]
categories = [*categories, categories[0]]

label_loc = np.linspace(start = 0, stop = 2*np.pi, num = len(scores))


plt.figure(figsize = (7,7))

ax = plt.subplot(polar = True) 

plt.xticks(label_loc, labels = categories, fontsize = 13)
ax.plot(label_loc, scores, label = 'Student1', linestyle = 'dashed', color = 'lightcoral')
ax.fill(label_loc, scores, color = 'lightcoral', alpha = 0.3)
ax.legend()


fig = plt.gcf()
fig.savefig("Chart.png") # 다각형 차트 생성 

document = Document() # 문서 생성

head = document.add_heading(level = 0)
head.add_run("목차").bold = True
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = document.add_paragraph()
para.add_run("\n1. IntelliKare란?\n").bold = True
para.add_run("\n2. 7가지 소검사 소개\n").bold = True
para.add_run("\n3. 종합 결과 분석\n").bold = True
para.add_run("\n4. 강점과 약점").bold = True
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

category = document.paragraphs[1].runs
for run in category:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(20)
    

document.add_page_break() # 1page

head = document.add_heading(level = 0)
head.add_run("IntelliKare란?").bold = True
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = document.add_paragraph()
para.add_run("\n IntelliKare는 VR을 활용한 웩슬러 기반의 아동 지능 검사입니다. 웩슬러 지능 검사는 신뢰성이 높은 검사로, 법적 의미를 갖는 유일한 지능 검사입니다. 검사 대상은 만 6세 ~ 16세이고, 5개의 지표 점수 산출을 통해 전체 지능을 확인합니다.\n")
para.add_run("\n IntelliKare는 기존 웩슬러 검사에서 측정하는 5가지 지표(언어 이해, 시공간, 유동추론, 작업 기억, 처리 속도)에 새로운 2가지 지표(반응 속도, 음악 지능)를 추가해서 총 7개의 영역을 검사합니다.\n")
para.add_run("\n 지능 검사 시스템을 구현할 때, 웩슬러 검사와 더불어 기존의 다른 지능 검사에서 제시한 문제 방식을 분석하고 최대한 반영하여 검사의 신뢰성을 높이고 VR 시스템에 적합한 검사 콘텐츠를 구현했습니다.\n")
para.add_run("\n VR 기술을 통해 만들어진 가상 현실 속에서 진행되는 스토리를 따라서 검사가 진행됩니다. 피검사자는 가상공간 내에 실제로 존재하는 것처럼 느끼고 자연스러운 상호작용을 통해 검사에 더욱 몰입할 수 있습니다. 그 안에 게임적인 요소들을 추가해서 동기 부여와 흥미 유발을 일으킵니다.\n")

content = document.paragraphs[4].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(11)

document.add_picture('vr.jpg', width = Inches(2.0))

document.add_page_break() # 2page

head = document.add_heading(level = 0)
head.add_run("7가지 소검사 소개").bold = True
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

introduction = document.add_paragraph()
introduction.add_run(" 피검사자는 스토리가 진행됨에 따라 아래의 소검사를 통해 언어이해, 시공간, 유동추론, 작업기억, 처리속도, 반응속도, 음악지능 각각의 지표가 측정됩니다.\n")
run = document.paragraphs[8].runs[0]
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
run.font.size = Pt(9)

para = document.add_paragraph()
para.add_run("1. 토막짜기:\n").bold = True
para.add_run(" 토막짜기 소검사에서 추상적인 시각 자극을 분석하고 종합하는 능력을 측정하기 위해서, 아동은 정해진 시간 내에 주어진 토막들을 이용해 제시된 모형과 똑같은 모양을 만들어야 합니다. 또한, 비언어적 추론 능력, 시지각 조직화 능력, 동시처리 능력, 시각-운동 협응 능력, 시각 자극의 전경과 배경을 분리하는 능력들을 측정할 수 있습니다.\n")
run = document.paragraphs[9].runs[0]
run.font.size = Pt(12)
run.underline = True

content = document.paragraphs[9].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    

para = document.add_paragraph()
para.add_run("2. 어휘:\n").bold = True
para.add_run(" 어휘 소검사에서 아동의 언어적 개념 표현 능력을 측정하기 위해, 아동은 제시된 그림을 보고 그 단어의 이름을 말해야 합니다.  결정지능, 어휘지식, 장기기억, 언어 유창성, 학습 능력, 청각, 추상적 사고 능력 등의 지적 능력들을 측정할 수 있습니다.\n")
run = document.paragraphs[10].runs[0]
run.font.size = Pt(12)
run.underline = True

content = document.paragraphs[10].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


para = document.add_paragraph()
para.add_run("3. 행렬추리:\n").bold = True
para.add_run(" 행렬추리 소검사에서 아동의 시지각 정보처리 증력과 추상적 추론 능력을 측정하기 위해서, 아동은 행렬이나 연속의 일부를 보고, 전체를 완성하는 보기를 찾아야 합니다. 추가적으로, 유동지능, 작업기억, 동시처리 능력, 부분-전체의 관계에 대한 이해, 그리고 세부사항에 대한 주의력을 측정할 수 있습니다.\n")
run = document.paragraphs[11].runs[0]
run.font.size = Pt(12)
run.underline = True

content = document.paragraphs[11].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

document.add_page_break() # 3page

head = document.add_heading(level = 0)
head.add_run("7가지 소검사 소개").bold = True
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = document.add_paragraph()
para.add_run("4. 동형찾기:\n").bold = True
para.add_run(" 동형찾기 소검사에서 아동의 시각적 변별 능력, 시각운동 협응 능력, 주의집중력, 의사결정 속도, 통제 및 조절 능력, 시각단기기억 등을 측정하기 위해선 아동은 제한시간 내에 장난감의 기호를 훑어보고 문제의 모양과 동일한 것을 찾아야 합니다. 또한, 시지각 조직화 능력, 유동지능, 계획 및 학습 능력과 연관이 있습니다.\n")
run = document.paragraphs[14].runs[0]
run.font.size = Pt(12)
run.underline = True

content = document.paragraphs[14].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    

para = document.add_paragraph()
para.add_run("5. 그림기억:\n").bold = True
para.add_run(" 그림기억 소검사에서 아동의 작업기억, 주의집중력, 시각처리 능력, 순서화 능력, 시각단기기억과 반응 억제 능력을 측정합니다. 아동은 제한시간 내에 제시된 그림과 그림들의 순서를 기억한 다음, 해당 그림을 순서대로 나열해야 됩니다.\n")
run = document.paragraphs[15].runs[0]
run.font.size = Pt(12)
run.underline = True

content = document.paragraphs[15].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


para = document.add_paragraph()
para.add_run("6. 반응성:\n").bold = True
para.add_run(" 반응성 소검사에서 아동의 반응 속도와 더불어 상지 민첩성 및 손 기민성을 측정합니다. 아동은 빠른 속도로 임의로 생성되는 몬스터들을 총을 쏴서 제거하는 게임의 형태로 검사가 진행됩니다.\n")
run = document.paragraphs[16].runs[0]
run.font.size = Pt(12)
run.underline = True

content = document.paragraphs[16].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

para = document.add_paragraph()
para.add_run("7. 리듬 맞추기:\n").bold = True
para.add_run(" 음악지능 소검사에서 아동의 음악적 패턴과 리듬, 음조 등을 파악하는 능력을 측정합니다. 아동은 주어진 소리를 듣고 똑같은 리듬과 박자로 따라 연주하며 검사가 진행됩니다.\n")
run = document.paragraphs[17].runs[0]
run.font.size = Pt(12)
run.underline = True

content = document.paragraphs[17].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

document.add_page_break() # 4page

head = document.add_heading(level = 0)
head.add_run("종합 결과 분석").bold = True
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = document.add_paragraph()
r = para.add_run()
r.add_picture("Test.png", width = Inches(4.0))
r.add_picture("Score.png", width = Inches(4.0))

para = document.add_paragraph()
para.add_run(" 동의 언어이해 지표(VCI)는 108로 ‘평균 수준’에 해당하며, 이는 백분위 70%로 100명 가운데 상위 30%에 해당된다. 아동의 지각추론 지표(PRI)는 100으로 ‘평균 수준’에 해당하며, 이는 백분위 50%로 100명 가운데 상위 50%에 해당된다.")
run = document.paragraphs[21].runs[0]
run.font.size = Pt(11)
run.underline = True

content = document.paragraphs[21].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

document.add_page_break() # 5page

head = document.add_heading(level = 0)
head.add_run("강점과 약점").bold = True
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_picture("Chart.png", width = Inches(4.5))

para = document.add_paragraph()
para.add_run(" 지각추론 지표 중 '공통 그림 찾기’소검사에서 가장 낮은 점수를 얻었는데, 이는 아동이 다른 지각추론 영역에 비해 추상화 및 범주적 추론 능력이 저하되어 있음을 나타낸다. 더불어 시각적 자극에 대한 추상적 추론 능력을 측정하는 '공통 그림 찾기’소검사에 비해 언어적 자극에 대한 추론능력을 측정하는‘공통성’소검사가 유의미하게 높게 나타났다는 점에서, 아동이 과제나 자극을 비언어적으로 부여받을 때보다 언어적으로 부여받았을 때 보다 쉽게 문제를 추론하고 해결해 나갈 수 있다고 유추할 수 있다.") 
run = document.paragraphs[25].runs[0]
run.font.size = Pt(11)
run.underline = True

content = document.paragraphs[25].runs
for run in content:
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

document.save("IntelliKare 결과지.docx")
